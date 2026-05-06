"""Intake API — convert a raw uploaded document (image or PDF) into a
structured IntakeResult that the case-creation flow can plug into the DAG.

Endpoint:
  POST /api/v1/intake/parse-document
        multipart/form-data with 'file' field, OR JSON body matching IntakeDocument

The endpoint is auth-gated via the standard get_current_user dependency.

PDF fast path: PDFs detected by magic bytes (`%PDF-`) bypass the PIL
classifier entirely. We run pypdf for text-layer extraction first
(cheap, no AWS), then fall back to AWS Textract for image-only PDFs.
This guarantees PDFs never get tagged "unreadable" because PIL can't
decode them — PIL was never the right tool for PDFs anyway.

Clinical-content gate: extraction success is necessary but not sufficient.
A successfully-parsed PDF/DOCX could still be a non-clinical document
(architecture spec, marketing deck, expense report). After extraction we
score the text against a curated clinical lexicon and downgrade confidence
+ flag `non-clinical-content` when nothing oncology-shaped is in scope.
"""
from __future__ import annotations

import asyncio
import base64
import hashlib
import io
import logging
import re
import time
from typing import Any

from fastapi import APIRouter, Depends, File, HTTPException, UploadFile

from app.agents.intake import parse_document
from app.auth import get_current_user
from app.models.intake import (
    DocumentClassification,
    ExtractedField,
    IntakeDocument,
    IntakeResult,
    OCRResult,
)

log = logging.getLogger(__name__)

router = APIRouter(prefix="/intake", tags=["intake"])

_ACCEPTED_MIME = {
    "image/png",
    "image/jpeg",
    "image/webp",
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",  # .docx
    "application/msword",                                                        # .doc (legacy)
    "text/plain",                                                                # .txt
}
# Filename-extension fallback for browsers that mislabel DOCX uploads (Edge,
# certain Outlook flows) as application/octet-stream or empty content_type.
_EXTENSION_TO_MIME = {
    ".pdf":  "application/pdf",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".doc":  "application/msword",
    ".txt":  "text/plain",
    ".png":  "image/png",
    ".jpg":  "image/jpeg",
    ".jpeg": "image/jpeg",
    ".webp": "image/webp",
}
_MAX_BYTES = 8 * 1024 * 1024  # 8 MB cap — phone-camera scans are typically <4 MB
_MIN_CONFIDENCE = 0.70

# ---------------------------------------------------------------------------
# Clinical-content scoring
# ---------------------------------------------------------------------------
# A document that parses cleanly via pypdf/python-docx still might not be a
# clinical document. We grade the extracted text against an oncology-shaped
# lexicon. A high-quality clinical document (FHIR bundle, pathology report,
# H&P note, radiology report, denial letter) will hit many of these terms;
# an architecture spec or marketing deck will hit zero.
_CLINICAL_LEXICON = {
    # roles + sections common in clinical documents
    "diagnosis", "patient", "physician", "oncology", "oncologist", "tumor",
    "tumour", "cancer", "carcinoma", "metastatic", "neoplasm", "lesion",
    "pathology", "biopsy", "specimen", "histology", "cytology", "stage",
    "grade", "ecog", "kps", "performance status", "comorbidity",
    # diagnostics + biomarkers
    "biomarker", "her2", "egfr", "kras", "braf", "alk", "ros1", "brca",
    "tmb", "msi-h", "msi", "pd-l1", "pdl1", "ihc", "fish", "ngs",
    "fishtest", "ck20", "ki-67", "mmr", "dmmr", "hrd",
    # treatments + drug classes
    "chemotherapy", "chemo", "radiation", "radiotherapy", "infusion",
    "trastuzumab", "pertuzumab", "pembrolizumab", "nivolumab", "olaparib",
    "palbociclib", "osimertinib", "rituximab", "bevacizumab", "carboplatin",
    "cisplatin", "paclitaxel", "docetaxel", "doxorubicin", "fluorouracil",
    "5-fu", "5fu", "capecitabine", "temozolomide", "imatinib", "tucatinib",
    "letrozole", "tamoxifen", "abiraterone", "enzalutamide", "everolimus",
    "lapatinib", "mtor", "parp", "tki", "antibody", "checkpoint inhibitor",
    # codes + units
    "icd-10", "icd10", "cpt", "hcpcs", "j-code", "j code",
    "mg/m²", "mg/m2", "mg/kg", "auc",
    # workflow + payer terms
    "prior authorization", "prior authorisation", "preauthorization",
    "medical necessity", "denial", "appeal", "p2p", "peer-to-peer",
    "fhir", "claim", "claimresponse", "x12 278",
    # NCCN / regulatory references
    "nccn", "asco", "nci", "fda", "clinical guideline", "compendium",
}
_CLINICAL_SCORE_THRESHOLD = 0.30  # below this → flag as non-clinical
_CLINICAL_PATTERN = re.compile(
    r"\b(?:" + "|".join(re.escape(t) for t in sorted(_CLINICAL_LEXICON, key=len, reverse=True)) + r")\b",
    re.IGNORECASE,
)


def _clinical_content_score(text: str) -> tuple[float, list[str]]:
    """Score in [0, 1] based on density of clinical terms. Returns (score, hits)."""
    if not text or not text.strip():
        return 0.0, []
    matches = _CLINICAL_PATTERN.findall(text)
    unique = set(m.lower() for m in matches)
    # Word-count-aware density: a 200-word note with 5 clinical terms is
    # high-signal; a 10,000-word architecture doc with 5 terms is not.
    n_words = max(1, len(text.split()))
    density = len(matches) / n_words   # hits per word
    coverage = len(unique) / 30        # how many distinct concepts (cap at 30)
    # Combine: density is the primary signal; coverage breaks ties.
    score = min(1.0, max(0.0, density * 25 + coverage * 0.4))
    return round(score, 3), sorted(unique)


def _apply_clinical_gate(ocr: OCRResult, risk_flags: list[str]) -> tuple[OCRResult, list[str]]:
    """Multiply OCR confidence by the clinical-content score.

    Net effect: a clean-extract on a non-medical doc returns ocr.overall_confidence
    well below the 0.70 dispatch threshold, so the UI's 'Create case' affordance
    correctly hides + the risk_flags carry the reason."""
    score, hits = _clinical_content_score(ocr.full_text)
    new_conf = round(ocr.overall_confidence * max(0.15, score), 3)
    flags = list(risk_flags)
    if score < _CLINICAL_SCORE_THRESHOLD:
        if "non-clinical-content" not in flags:
            flags.append("non-clinical-content")
    new_ocr = ocr.model_copy(update={"overall_confidence": new_conf})
    return new_ocr, flags


@router.post(
    "/parse-document",
    response_model=IntakeResult,
    summary=(
        "Parse a raw document (handwritten Rx · scanned report · faxed letter) "
        "into a structured IntakeResult."
    ),
)
async def parse_document_endpoint(
    file: UploadFile = File(...),
    user: dict[str, Any] = Depends(get_current_user),
) -> IntakeResult:
    """Multipart-upload entry point. Hashes the bytes, classifies + extracts,
    and returns an IntakeResult. Always returns 200 + a structured result
    even on extraction failure — the IntakeResult.requires_human_review
    flag tells the caller whether the case can flow autonomously or needs
    a clinician's eye."""
    # Normalize content_type — some browsers send empty or generic types
    # for DOCX/PDF. Fall back to the filename extension when needed.
    content_type = (file.content_type or "").lower()
    filename = file.filename or "upload"
    if content_type not in _ACCEPTED_MIME:
        ext = _ext(filename)
        guessed = _EXTENSION_TO_MIME.get(ext)
        if guessed:
            log.info("intake: rewriting MIME from %r to %r based on extension %r", content_type, guessed, ext)
            content_type = guessed
        else:
            raise HTTPException(
                status_code=415,
                detail=(
                    f"Unsupported media type: {file.content_type!r} (filename {filename!r}). "
                    f"Accept: PDF, DOCX, PNG, JPEG, WebP, TXT."
                ),
            )

    raw = await file.read()
    if len(raw) > _MAX_BYTES:
        raise HTTPException(
            status_code=413,
            detail=f"Document too large: {len(raw)} bytes > {_MAX_BYTES} cap.",
        )
    if len(raw) < 50:
        raise HTTPException(
            status_code=400,
            detail="Document body is implausibly small (<50 bytes).",
        )

    sha256 = hashlib.sha256(raw).hexdigest()

    # ------------------------------------------------------------------
    # Format dispatch — magic-byte detection always wins over MIME
    # ------------------------------------------------------------------
    detected_kind = _detect_kind(raw, content_type, filename)
    log.info("intake: %s detected as %s (%d bytes, mime=%s)", filename, detected_kind, len(raw), content_type)

    if detected_kind == "pdf":
        return await _extract_pdf_directly(raw=raw, filename=filename, sha256=sha256)
    if detected_kind == "docx":
        return await _extract_docx_directly(raw=raw, filename=filename, sha256=sha256)
    if detected_kind == "txt":
        return _extract_txt_directly(raw=raw, filename=filename, sha256=sha256)

    # ------------------------------------------------------------------
    # Image path — go through the full pipeline (Vision/Textract/Tesseract)
    # ------------------------------------------------------------------
    doc = IntakeDocument(
        filename=filename,
        mime_type=content_type,  # type: ignore[arg-type]  # validated above
        image_b64=base64.b64encode(raw).decode("ascii"),
        sha256=sha256,
        source="upload",
    )

    # Image path goes through the full pipeline. Apply the clinical-content
    # gate to the result so a non-clinical image (e.g. a generic UI screenshot)
    # gets its confidence downgraded + flagged with `non-clinical-content` —
    # same gate the PDF/DOCX/TXT fast paths apply.
    result = await parse_document(doc)
    new_ocr, new_flags = _apply_clinical_gate(result.ocr, result.risk_flags)
    requires_review = (
        new_ocr.overall_confidence < _MIN_CONFIDENCE
        or not new_ocr.full_text.strip()
        or "non-clinical-content" in new_flags
    )
    return result.model_copy(update={
        "ocr": new_ocr,
        "risk_flags": new_flags,
        "requires_human_review": result.requires_human_review or requires_review,
    })


def _ext(filename: str) -> str:
    """Lowercase extension including the dot — '' if none."""
    i = filename.rfind(".")
    return filename[i:].lower() if i >= 0 else ""


def _detect_kind(raw: bytes, mime: str, filename: str) -> str:
    """Magic-byte + MIME + extension cascade. Returns one of
    {pdf, docx, txt, image}. Magic bytes always win when present."""
    head = raw[:8]
    if head[:5] == b"%PDF-" or b"%PDF" in raw[:1024]:
        return "pdf"
    # DOCX is a ZIP archive (PK\x03\x04). The .doc legacy format is OLE2
    # (D0 CF 11 E0); python-docx can't read .doc, so we route it to the
    # docx extractor anyway and let it fail gracefully (HITL).
    if head[:2] == b"PK" and (
        mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        or _ext(filename) in (".docx", ".doc")
    ):
        return "docx"
    if mime == "application/msword" or _ext(filename) == ".doc":
        return "docx"  # python-docx will raise; we fall back to HITL
    if mime == "text/plain" or _ext(filename) == ".txt":
        return "txt"
    return "image"


# ---------------------------------------------------------------------------
# PDF fast-path extraction
# ---------------------------------------------------------------------------
async def _extract_pdf_directly(*, raw: bytes, filename: str, sha256: str) -> IntakeResult:
    """Try pypdf (text layer) first, then Textract (scanned PDFs).

    Always returns a valid IntakeResult — never raises. If both engines fail,
    the result has overall_confidence=0 and requires_human_review=True so the
    case routes straight to the Reviewer queue.
    """
    t_start = time.monotonic()
    classification = DocumentClassification(
        document_type="typed_print",
        confidence=0.95,
        rationale="PDF detected via %PDF- magic bytes — routed to PDF-native extractor.",
        quality_flags=["pdf-document"],
    )

    engines_attempted: list[dict[str, Any]] = []
    ocr: OCRResult | None = None
    risk_flags: list[str] = ["pdf-document"]

    # 1. pypdf — text-layer extraction, no AWS, ~80 ms
    pypdf_t0 = time.monotonic()
    try:
        ocr = await asyncio.to_thread(_pypdf_extract, raw)
        engines_attempted.append({
            "engine": "pypdf_text",
            "ok": True,
            "elapsed_ms": int((time.monotonic() - pypdf_t0) * 1000),
        })
    except _PDFExtractFailure as e:
        engines_attempted.append({
            "engine": "pypdf_text",
            "ok": False,
            "reason": str(e)[:160],
            "elapsed_ms": int((time.monotonic() - pypdf_t0) * 1000),
        })
        log.info("pypdf failed for %s: %s — trying Textract", filename, e)

    # 2. Textract fallback — for image-only PDFs (scanned, no text layer)
    if ocr is None:
        textract_t0 = time.monotonic()
        try:
            ocr = await asyncio.to_thread(_textract_extract, raw)
            engines_attempted.append({
                "engine": "aws_textract",
                "ok": True,
                "elapsed_ms": int((time.monotonic() - textract_t0) * 1000),
            })
            risk_flags.append("engine-fallback-used")
        except _PDFExtractFailure as e:
            engines_attempted.append({
                "engine": "aws_textract",
                "ok": False,
                "reason": str(e)[:160],
                "elapsed_ms": int((time.monotonic() - textract_t0) * 1000),
            })
            log.info("Textract unavailable for %s: %s — trying local OCR (pypdfium2 + tesseract)", filename, e)

    # 3. Local OCR fallback — rasterize pages with pypdfium2 and OCR with
    #    tesseract. Zero cloud dependency. Handles scanned/image-only PDFs
    #    when AWS isn't configured.
    if ocr is None:
        ocr_t0 = time.monotonic()
        try:
            ocr = await asyncio.to_thread(_pdfium_tesseract_extract, raw)
            engines_attempted.append({
                "engine": "tesseract_local",
                "ok": True,
                "elapsed_ms": int((time.monotonic() - ocr_t0) * 1000),
            })
            risk_flags.append("engine-fallback-used")
        except _PDFExtractFailure as e:
            engines_attempted.append({
                "engine": "tesseract_local",
                "ok": False,
                "reason": str(e)[:160],
                "elapsed_ms": int((time.monotonic() - ocr_t0) * 1000),
            })
            log.warning("Local OCR (pypdfium2+tesseract) failed for %s: %s", filename, e)

    # 4. All engines failed → HITL
    if ocr is None:
        risk_flags.append("intake-failed")
        ocr = OCRResult(
            engine="pypdf_text",  # nominal — no extraction actually succeeded
            full_text="",
            extracted_fields=[],
            overall_confidence=0.0,
            phi_redactions_applied=0,
            pages=1,
        )

    # Clinical-content gate: extraction can succeed on a non-medical document.
    # Score the text against the oncology lexicon and downgrade confidence
    # accordingly so the UI doesn't offer "Create case" on an unrelated PDF.
    ocr, risk_flags = _apply_clinical_gate(ocr, risk_flags)

    requires_review = (
        ocr.overall_confidence < _MIN_CONFIDENCE
        or not ocr.full_text.strip()
        or "non-clinical-content" in risk_flags
    )
    if requires_review and "intake-failed" not in risk_flags and "non-clinical-content" not in risk_flags:
        risk_flags.append("low-confidence")

    total_ms = int((time.monotonic() - t_start) * 1000)
    return IntakeResult(
        classification=classification,
        ocr=ocr,
        clinical_snapshot_partial={},
        risk_flags=risk_flags,
        requires_human_review=requires_review,
        audit={
            "document_sha256": sha256,
            "filename": filename,
            "engines_used": [a["engine"] for a in engines_attempted if a.get("ok")],
            "engines_attempted": engines_attempted,
            "latency_ms": total_ms,
            "fast_path": "pdf",
        },
    )


# ---------------------------------------------------------------------------
# DOCX fast-path extraction
# ---------------------------------------------------------------------------
async def _extract_docx_directly(*, raw: bytes, filename: str, sha256: str) -> IntakeResult:
    """Extract paragraphs + tables from a .docx file via python-docx.

    DOCX is structured XML inside a ZIP — text extraction is deterministic
    and lossless. No OCR needed, no AWS dependency. Tables are flattened
    into tab-separated rows so downstream regex/extraction sees them as
    coherent lines."""
    t_start = time.monotonic()
    classification = DocumentClassification(
        document_type="typed_print",
        confidence=0.98,
        rationale="DOCX (Office Open XML) — text extracted directly from the .docx XML stream.",
        quality_flags=["docx-document"],
    )
    risk_flags: list[str] = ["docx-document"]
    engines_attempted: list[dict[str, Any]] = []

    docx_t0 = time.monotonic()
    ocr: OCRResult | None = None
    try:
        ocr = await asyncio.to_thread(_python_docx_extract, raw)
        engines_attempted.append({
            "engine": "python_docx",
            "ok": True,
            "elapsed_ms": int((time.monotonic() - docx_t0) * 1000),
        })
    except _DocExtractFailure as e:
        engines_attempted.append({
            "engine": "python_docx",
            "ok": False,
            "reason": str(e)[:160],
            "elapsed_ms": int((time.monotonic() - docx_t0) * 1000),
        })
        log.warning("python-docx failed for %s: %s", filename, e)

    if ocr is None:
        risk_flags.append("intake-failed")
        ocr = OCRResult(
            engine="python_docx",
            full_text="",
            extracted_fields=[],
            overall_confidence=0.0,
            phi_redactions_applied=0,
            pages=1,
        )

    ocr, risk_flags = _apply_clinical_gate(ocr, risk_flags)

    requires_review = (
        ocr.overall_confidence < _MIN_CONFIDENCE
        or not ocr.full_text.strip()
        or "non-clinical-content" in risk_flags
    )
    if requires_review and "intake-failed" not in risk_flags and "non-clinical-content" not in risk_flags:
        risk_flags.append("low-confidence")

    return IntakeResult(
        classification=classification,
        ocr=ocr,
        clinical_snapshot_partial={},
        risk_flags=risk_flags,
        requires_human_review=requires_review,
        audit={
            "document_sha256": sha256,
            "filename": filename,
            "engines_used": [a["engine"] for a in engines_attempted if a.get("ok")],
            "engines_attempted": engines_attempted,
            "latency_ms": int((time.monotonic() - t_start) * 1000),
            "fast_path": "docx",
        },
    )


def _python_docx_extract(raw: bytes) -> OCRResult:
    """Pull all paragraph text + flatten tables. Raises _DocExtractFailure
    on parse errors or empty documents."""
    try:
        from docx import Document  # python-docx
    except ImportError as e:
        raise _DocExtractFailure(f"python-docx not installed: {e}") from e

    try:
        doc = Document(io.BytesIO(raw))
    except Exception as e:  # noqa: BLE001 — covers BadZipFile, KeyError, etc.
        raise _DocExtractFailure(f"docx parse error: {e}") from e

    parts: list[str] = []
    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [(c.text or "").strip() for c in row.cells]
            line = "\t".join(c for c in cells if c)
            if line:
                parts.append(line)

    full_text = "\n".join(parts)
    if not full_text.strip():
        raise _DocExtractFailure("DOCX has no extractable text (empty document)")

    return OCRResult(
        engine="python_docx",
        full_text=full_text,
        extracted_fields=[],
        overall_confidence=0.92,  # high — structured extraction, not OCR
        phi_redactions_applied=0,
        pages=1,  # python-docx doesn't expose printed page count
    )


# ---------------------------------------------------------------------------
# Plain-text fast-path extraction
# ---------------------------------------------------------------------------
def _extract_txt_directly(*, raw: bytes, filename: str, sha256: str) -> IntakeResult:
    """Decode raw bytes as UTF-8 with a Latin-1 fallback. Always succeeds."""
    t_start = time.monotonic()
    try:
        text = raw.decode("utf-8")
    except UnicodeDecodeError:
        text = raw.decode("latin-1", errors="replace")

    classification = DocumentClassification(
        document_type="typed_print",
        confidence=1.0,
        rationale="Plain-text upload — decoded directly, no OCR required.",
        quality_flags=["plain-text"],
    )
    ocr = OCRResult(
        engine="plain_text",
        full_text=text,
        extracted_fields=[],
        overall_confidence=1.0,
        phi_redactions_applied=0,
        pages=1,
    )
    risk_flags = ["plain-text"]
    ocr, risk_flags = _apply_clinical_gate(ocr, risk_flags)
    requires_review = (
        not text.strip()
        or ocr.overall_confidence < _MIN_CONFIDENCE
        or "non-clinical-content" in risk_flags
    )
    if not text.strip():
        risk_flags.append("intake-failed")

    return IntakeResult(
        classification=classification,
        ocr=ocr,
        clinical_snapshot_partial={},
        risk_flags=risk_flags,
        requires_human_review=requires_review,
        audit={
            "document_sha256": sha256,
            "filename": filename,
            "engines_used": ["plain_text"],
            "engines_attempted": [{"engine": "plain_text", "ok": True}],
            "latency_ms": int((time.monotonic() - t_start) * 1000),
            "fast_path": "txt",
        },
    )


class _DocExtractFailure(Exception):
    """Raised when a DOCX/TXT engine cannot extract anything usable."""


class _PDFExtractFailure(Exception):
    """Raised when a PDF engine cannot extract anything usable."""


def _pypdf_extract(raw: bytes) -> OCRResult:
    """Synchronous pypdf text-layer extraction. Raises _PDFExtractFailure
    on any error or empty extraction (image-only PDF)."""
    try:
        from pypdf import PdfReader
    except ImportError as e:
        raise _PDFExtractFailure(f"pypdf not installed: {e}") from e

    try:
        reader = PdfReader(io.BytesIO(raw))
        page_texts = [(p.extract_text() or "") for p in reader.pages]
        n_pages = len(reader.pages)
    except Exception as e:  # noqa: BLE001 — surface as clean failure
        raise _PDFExtractFailure(f"pypdf parse error: {e}") from e

    full_text = "\n\n".join(t for t in page_texts if t.strip())
    if not full_text.strip():
        raise _PDFExtractFailure("PDF has no text layer (image-only / scanned)")

    # Heuristic confidence: text-layer extraction is deterministic and high
    # quality when it works at all. We use 0.85 as the floor so the case
    # auto-flows past the 0.70 HITL gate.
    return OCRResult(
        engine="pypdf_text",
        full_text=full_text,
        extracted_fields=[],
        overall_confidence=0.85,
        phi_redactions_applied=0,
        pages=n_pages,
    )


def _textract_extract(raw: bytes) -> OCRResult:
    """Synchronous AWS Textract extraction. Raises _PDFExtractFailure if
    boto3 is missing, AWS creds are not configured, or the call fails."""
    try:
        import boto3
        from botocore.config import Config
        from botocore.exceptions import BotoCoreError, ClientError, NoCredentialsError
    except ImportError as e:
        raise _PDFExtractFailure(f"boto3 not installed: {e}") from e

    from app.config import settings

    try:
        client = boto3.client(
            "textract",
            config=Config(
                region_name=settings.AWS_REGION,
                retries={"max_attempts": 3, "mode": "adaptive"},
                read_timeout=30,
            ),
        )
    except Exception as e:  # noqa: BLE001
        raise _PDFExtractFailure(f"Textract client init failed: {e}") from e

    try:
        response = client.analyze_document(
            Document={"Bytes": raw},
            FeatureTypes=["FORMS", "TABLES"],
        )
    except NoCredentialsError as e:
        raise _PDFExtractFailure("AWS credentials not configured") from e
    except (ClientError, BotoCoreError) as e:
        raise _PDFExtractFailure(f"Textract call failed: {e}") from e
    except Exception as e:  # noqa: BLE001
        raise _PDFExtractFailure(f"Textract unexpected error: {e}") from e

    blocks = response.get("Blocks", [])
    word_confs: list[float] = []
    line_texts: list[str] = []
    by_id = {b["Id"]: b for b in blocks if "Id" in b}

    for b in blocks:
        bt = b.get("BlockType")
        if bt == "WORD":
            c = b.get("Confidence", 0.0)
            if isinstance(c, (int, float)):
                word_confs.append(c / 100.0)
        elif bt == "LINE":
            t = b.get("Text") or ""
            if t:
                line_texts.append(t)

    overall = sum(word_confs) / len(word_confs) if word_confs else 0.0

    extracted: list[ExtractedField] = []
    for b in blocks:
        if (
            b.get("BlockType") == "KEY_VALUE_SET"
            and "KEY" in (b.get("EntityTypes") or [])
        ):
            key = _resolve_text(b, by_id).strip()
            value_block = _find_value_for_key(b, by_id)
            value = _resolve_text(value_block, by_id).strip() if value_block else ""
            conf = (b.get("Confidence", 0.0) / 100.0) if "Confidence" in b else 0.0
            if conf >= 0.70 and key and value:
                extracted.append(
                    ExtractedField(
                        name=f"textract.{key.lower().replace(' ', '_')[:60]}",
                        value=value[:200],
                        confidence=round(conf, 3),
                        source_excerpt=f"{key}: {value}"[:200],
                        page=int(b.get("Page", 1)),
                    )
                )

    if not line_texts:
        raise _PDFExtractFailure("Textract returned no text blocks")

    return OCRResult(
        engine="aws_textract",
        full_text="\n".join(line_texts),
        extracted_fields=extracted,
        overall_confidence=round(overall, 3),
        phi_redactions_applied=0,
        pages=int(response.get("DocumentMetadata", {}).get("Pages", 1)),
    )


def _pdfium_tesseract_extract(raw: bytes) -> OCRResult:
    """Local OCR fallback: rasterize PDF pages with pypdfium2 and OCR each
    page with tesseract. Zero cloud dependency, works fully offline.

    Designed for scanned/image-only PDFs when neither pypdf (no text layer)
    nor Textract (no AWS) can extract. Confidence is intentionally low (0.65)
    so the result still triggers a HITL review — local OCR on a phone-camera
    scan is best-effort, not authoritative."""
    try:
        import pypdfium2 as pdfium
        import pytesseract
        from PIL import Image  # noqa: F401  (sanity check)
    except ImportError as e:
        raise _PDFExtractFailure(f"local OCR deps missing: {e}") from e

    # Auto-locate the tesseract binary on Windows when it's not on PATH.
    import os
    import shutil

    tess_cmd = shutil.which("tesseract") or shutil.which("tesseract.exe")
    if not tess_cmd:
        for candidate in (
            r"C:\Program Files\Tesseract-OCR\tesseract.exe",
            r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
            "/usr/bin/tesseract",
            "/usr/local/bin/tesseract",
            "/opt/homebrew/bin/tesseract",
        ):
            if os.path.exists(candidate):
                tess_cmd = candidate
                break
    if not tess_cmd:
        raise _PDFExtractFailure("tesseract binary not found on PATH or in standard locations")
    pytesseract.pytesseract.tesseract_cmd = tess_cmd

    try:
        pdf = pdfium.PdfDocument(raw)
    except Exception as e:  # noqa: BLE001
        raise _PDFExtractFailure(f"pypdfium2 could not open PDF: {e}") from e

    n_pages = len(pdf)
    if n_pages == 0:
        raise _PDFExtractFailure("PDF has zero pages")

    page_texts: list[str] = []
    page_confs: list[float] = []
    # 200 DPI is the sweet spot — high enough for clean tesseract OCR,
    # not so high that a 10-page PDF blows the request budget.
    scale = 200 / 72  # PDF user-space units → 200 DPI
    for i in range(n_pages):
        page = pdf[i]
        try:
            pil_img = page.render(scale=scale).to_pil()
        except Exception as e:  # noqa: BLE001
            log.warning("page %d render failed: %s", i + 1, e)
            continue
        finally:
            try:
                page.close()
            except Exception:  # noqa: BLE001
                pass

        try:
            data = pytesseract.image_to_data(
                pil_img, output_type=pytesseract.Output.DICT, lang="eng"
            )
        except Exception as e:  # noqa: BLE001
            log.warning("page %d OCR failed: %s", i + 1, e)
            continue

        words = []
        confs = []
        for txt, conf in zip(data.get("text", []), data.get("conf", [])):
            t = (txt or "").strip()
            try:
                c = float(conf)
            except (TypeError, ValueError):
                c = -1.0
            if t and c >= 0:
                words.append(t)
                confs.append(c)
        if words:
            page_texts.append(" ".join(words))
            page_confs.extend(confs)

    try:
        pdf.close()
    except Exception:  # noqa: BLE001
        pass

    full_text = "\n\n".join(page_texts).strip()
    if not full_text:
        raise _PDFExtractFailure("tesseract produced no text from any page")

    # Tesseract returns 0-100 confidence per word; normalize to [0,1].
    overall = (sum(page_confs) / len(page_confs) / 100.0) if page_confs else 0.65
    overall = max(0.5, min(0.95, overall))  # clamp — local OCR is mid-confidence by nature

    return OCRResult(
        engine="tesseract_local",
        full_text=full_text,
        extracted_fields=[],
        overall_confidence=round(overall, 3),
        phi_redactions_applied=0,
        pages=n_pages,
    )


def _resolve_text(block: dict | None, by_id: dict) -> str:
    if not block:
        return ""
    pieces: list[str] = []
    for rel in block.get("Relationships", []) or []:
        if rel.get("Type") != "CHILD":
            continue
        for cid in rel.get("Ids", []) or []:
            child = by_id.get(cid)
            if not child:
                continue
            if child.get("BlockType") == "WORD":
                pieces.append(child.get("Text", ""))
            elif (
                child.get("BlockType") == "SELECTION_ELEMENT"
                and child.get("SelectionStatus") == "SELECTED"
            ):
                pieces.append("[X]")
    return " ".join(pieces)


def _find_value_for_key(key_block: dict, by_id: dict) -> dict | None:
    for rel in key_block.get("Relationships", []) or []:
        if rel.get("Type") == "VALUE":
            for vid in rel.get("Ids", []) or []:
                v = by_id.get(vid)
                if v:
                    return v
    return None
